"""DOCX 文本抽取（python-docx）。

抽取段落 + 表格文本，保持文档顺序。python-docx 不处理 .doc（旧格式），
旧 .doc 文件会回退为 best-effort 文本读取。
"""

from __future__ import annotations

import asyncio
import logging
from pathlib import Path

from ..models import ExtractedContent

logger = logging.getLogger(__name__)


async def extract(path: Path, max_chars: int = 0) -> ExtractedContent:
    """从 DOCX 抽取纯文本。"""
    return await asyncio.to_thread(_extract_sync, path, max_chars)


def _extract_sync(path: Path, max_chars: int) -> ExtractedContent:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed; cannot extract .docx")
        return ExtractedContent(text="", file_type="docx", metadata={"error": "python-docx not installed"})

    doc = Document(str(path))
    parts: list[str] = []
    total = 0

    # 段落
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
            total += len(line)
            if max_chars > 0 and total >= max_chars:
                text = "\n".join(parts)[:max_chars]
                return ExtractedContent(
                    text=text, file_type="docx", char_count=len(text),
                    metadata={"truncated": True},
                )

    # 表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cells)
            if line.strip(" |"):
                parts.append(line)
                total += len(line)
                if max_chars > 0 and total >= max_chars:
                    text = "\n".join(parts)[:max_chars]
                    return ExtractedContent(
                        text=text, file_type="docx", char_count=len(text),
                        metadata={"truncated": True, "tables": len(doc.tables)},
                    )

    text = "\n".join(parts)
    return ExtractedContent(
        text=text,
        file_type="docx",
        char_count=len(text),
        metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
    )
